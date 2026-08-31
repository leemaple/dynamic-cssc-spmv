#!/usr/bin/env python3
"""Build the Route C English manuscript and Chinese technical note as DOCX.

Pandoc performs Markdown/citation/math conversion.  This script then applies a
deterministic Word-native style layer, fixed table geometry, running furniture,
and accessibility metadata without replacing Pandoc's editable OMML equations.
"""

from __future__ import annotations

import argparse
import shutil
import subprocess
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph

REPO = Path(__file__).resolve().parents[1]
PAPER = REPO / "docs" / "paper"
EN_MD = PAPER / "manuscript-draft.md"
EN_DOCX = PAPER / "manuscript-draft.docx"
ZH_MD = PAPER / "paper-idea-detailed-zh.md"
ZH_DOCX = PAPER / "paper-idea-detailed-zh.docx"
BIB = PAPER / "references.bib"

EN_SHORT = "VERSION-BOUND MUTABLE CSSC SPMV • ROUTE C WORKING MANUSCRIPT"
ZH_SHORT = "DYNAMIC CSSC SPMV • ROUTE C 技术说明"
# Use a system-wide Unicode TTF rather than a user-profile font or TTC.  The
# renderer gives LibreOffice an isolated HOME, so per-user fonts are invisible.
ZH_FONT = "Arial Unicode MS"
EN_CORE_TITLE = (
    "Version-Bound Maintenance for Mutable Homomorphic Sparse Matrix–Vector "
    "Multiplication: A Fail-Closed Evaluation Boundary"
)

INK = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5B6573"
LIGHT = "F4F6F9"
GRID = "CBD3DC"


def _run(command: list[str]) -> None:
    subprocess.run(command, cwd=REPO, check=True)


def _pandoc(source: Path, target: Path, *, reference: Path | None, cite: bool) -> None:
    command = [
        "pandoc",
        str(source),
        "--from=markdown+tex_math_dollars+implicit_figures",
        "--to=docx",
        f"--resource-path={PAPER}",
        "--standalone",
        "--output",
        str(target),
    ]
    if reference is not None:
        command.extend(["--reference-doc", str(reference)])
    if cite:
        command.extend(
            [
                "--citeproc",
                f"--bibliography={BIB}",
                "--metadata=link-citations:true",
            ]
        )
    else:
        command.append("--number-sections")
    _run(command)


def _ensure(parent, tag: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def _set_font_element(rpr, name: str) -> None:
    fonts = _ensure(rpr, "w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        fonts.set(qn(attr), name)
    fonts.set(qn("w:hint"), "eastAsia")


def _set_style_font(style, name: str, size: float, *, color: str | None = None, bold=None) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    if color:
        style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        style.font.bold = bold
    _set_font_element(style.element.get_or_add_rPr(), name)


def _set_run_font(run, name: str, size: float | None = None) -> None:
    run.font.name = name
    if size is not None:
        run.font.size = Pt(size)
    _set_font_element(run._element.get_or_add_rPr(), name)


def _set_paragraph_tokens(
    style,
    *,
    before: float,
    after: float,
    line: float,
    justify=False,
) -> None:
    fmt = style.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = line
    if justify:
        fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def _get_or_add_style(doc: Document, name: str, base: str = "Normal"):
    try:
        return doc.styles[name]
    except KeyError:
        style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles[base]
        return style


def _paragraph_border_and_fill(paragraph, *, fill: str, left: str) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    shd = _ensure(ppr, "w:shd")
    shd.set(qn("w:fill"), fill)
    borders = _ensure(ppr, "w:pBdr")
    edge = _ensure(borders, "w:left")
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), "18")
    edge.set(qn("w:space"), "8")
    edge.set(qn("w:color"), left)


def _configure_english_styles(doc: Document) -> None:
    font = "Calibri"
    for name in ("Normal", "Body Text", "First Paragraph"):
        try:
            style = doc.styles[name]
        except KeyError:
            continue
        _set_style_font(style, font, 11, color="111827")
        _set_paragraph_tokens(style, before=0, after=8, line=1.333, justify=True)

    title = doc.styles["Title"]
    _set_style_font(title, font, 21, color=INK, bold=True)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(10)
    title.paragraph_format.keep_with_next = True

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[name]
        _set_style_font(style, font, size, color=color, bold=True)
        _set_paragraph_tokens(style, before=before, after=after, line=1.0)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for name in ("Block Text", "Quote"):
        try:
            style = doc.styles[name]
        except KeyError:
            continue
        _set_style_font(style, font, 10, color="374151")
        _set_paragraph_tokens(style, before=4, after=7, line=1.25)
        style.paragraph_format.left_indent = Inches(0.18)
        style.paragraph_format.right_indent = Inches(0.08)

    caption = doc.styles["Caption"]
    _set_style_font(caption, font, 9, color=MUTED)
    _set_paragraph_tokens(caption, before=5, after=8, line=1.15)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.keep_together = True

    for name in ("List Bullet", "List Number", "List Continue"):
        try:
            style = doc.styles[name]
        except KeyError:
            continue
        _set_style_font(style, font, 11, color="111827")
        _set_paragraph_tokens(style, before=0, after=4, line=1.208)

    try:
        bibliography = doc.styles["Bibliography"]
    except KeyError:
        bibliography = _get_or_add_style(doc, "Bibliography")
    _set_style_font(bibliography, font, 9.5, color="1F2937")
    _set_paragraph_tokens(bibliography, before=0, after=4, line=1.15)
    bibliography.paragraph_format.left_indent = Inches(0.25)
    bibliography.paragraph_format.first_line_indent = Inches(-0.25)

    kicker = _get_or_add_style(doc, "Route C Kicker")
    _set_style_font(kicker, font, 9, color=BLUE, bold=True)
    _set_paragraph_tokens(kicker, before=0, after=5, line=1.0)
    kicker.paragraph_format.keep_with_next = True


def _promote_english_headings_and_add_kicker(doc: Document) -> None:
    title_paragraph = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(
            "Version-Bound Maintenance for Mutable Homomorphic Sparse Matrix"
        ):
            title_paragraph = paragraph
            break
    if title_paragraph is None:
        raise RuntimeError("English manuscript title paragraph was not found")

    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name
        if paragraph.text.strip().startswith(
            "Version-Bound Maintenance for Mutable Homomorphic Sparse Matrix"
        ):
            paragraph.style = doc.styles["Title"]
        elif style_name == "Heading 2":
            paragraph.style = doc.styles["Heading 1"]
        elif style_name == "Heading 3":
            paragraph.style = doc.styles["Heading 2"]
        elif style_name == "Heading 4":
            paragraph.style = doc.styles["Heading 3"]

    kicker = title_paragraph.insert_paragraph_before(
        "ROUTE C • METHODS AND EVIDENCE-BOUNDARY WORKING MANUSCRIPT"
    )
    kicker.style = doc.styles["Route C Kicker"]


def _configure_page(section, *, letter: bool) -> None:
    if letter:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True


def _append_page_field(paragraph) -> None:
    run = paragraph.add_run("Page ")
    _set_run_font(run, "Calibri", 8.5)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)


def _set_header_footer(section, label: str, *, font: str, first_label: str | None = None) -> None:
    header_labels = (
        (section.header, label),
        (section.first_page_header, label if first_label is None else first_label),
    )
    for header, header_label in header_labels:
        paragraph = header.paragraphs[0]
        paragraph.text = header_label
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            _set_run_font(run, font, 8.5)
            run.font.color.rgb = RGBColor.from_string(MUTED)
            run.bold = True

    for footer in (section.footer, section.first_page_footer):
        paragraph = footer.paragraphs[0]
        paragraph.clear()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph.paragraph_format.space_before = Pt(0)
        _append_page_field(paragraph)
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor.from_string(MUTED)


def _set_cell_margins(cell, *, top=80, start=120, bottom=80, end=120) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    mar = _ensure(tcpr, "w:tcMar")
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = _ensure(mar, f"w:{edge}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _shade_cell(cell, fill: str) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    shd = _ensure(tcpr, "w:shd")
    shd.set(qn("w:fill"), fill)


def _set_cell_width(cell, width: int) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    tcw = _ensure(tcpr, "w:tcW")
    tcw.set(qn("w:w"), str(width))
    tcw.set(qn("w:type"), "dxa")


def _table_widths(column_count: int) -> list[int]:
    patterns = {
        1: [9360],
        2: [2700, 6660],
        3: [2050, 3110, 4200],
        4: [1600, 2200, 2800, 2760],
    }
    if column_count in patterns:
        return patterns[column_count]
    base, remainder = divmod(9360, column_count)
    return [base + (1 if index < remainder else 0) for index in range(column_count)]


def _set_table_borders(table) -> None:
    tblpr = table._tbl.tblPr
    borders = _ensure(tblpr, "w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = _ensure(borders, f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), GRID)


def _format_tables(doc: Document, *, font: str) -> None:
    for table in doc.tables:
        column_count = len(table.columns)
        compact = len(table.rows) > 8
        widths = _table_widths(column_count)
        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        tblpr = table._tbl.tblPr
        tblw = _ensure(tblpr, "w:tblW")
        tblw.set(qn("w:w"), "9360")
        tblw.set(qn("w:type"), "dxa")
        tblind = _ensure(tblpr, "w:tblInd")
        tblind.set(qn("w:w"), "120")
        tblind.set(qn("w:type"), "dxa")
        layout = _ensure(tblpr, "w:tblLayout")
        layout.set(qn("w:type"), "fixed")
        _set_table_borders(table)

        grid = table._tbl.tblGrid
        for child in list(grid):
            grid.remove(child)
        for width in widths:
            column = OxmlElement("w:gridCol")
            column.set(qn("w:w"), str(width))
            grid.append(column)

        for row_index, row in enumerate(table.rows):
            trpr = row._tr.get_or_add_trPr()
            cant_split = _ensure(trpr, "w:cantSplit")
            cant_split.set(qn("w:val"), "1")
            if row_index == 0:
                header = _ensure(trpr, "w:tblHeader")
                header.set(qn("w:val"), "1")
            for column_index, cell in enumerate(row.cells):
                _set_cell_width(cell, widths[min(column_index, len(widths) - 1)])
                if compact:
                    _set_cell_margins(cell, top=45, start=100, bottom=45, end=100)
                else:
                    _set_cell_margins(cell)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                if row_index == 0:
                    _shade_cell(cell, LIGHT)
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(1 if compact else 3)
                    paragraph.paragraph_format.line_spacing = 1.02 if compact else 1.08
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        _set_run_font(run, font, 8.4 if compact else 9.2)
                        if row_index == 0:
                            run.bold = True


def _space_paragraphs_after_tables(doc: Document) -> None:
    """Keep body text visually separate from the preceding table border."""
    for table in doc.tables:
        next_element = table._tbl.getnext()
        if next_element is not None and next_element.tag == qn("w:p"):
            Paragraph(next_element, table._parent).paragraph_format.space_before = Pt(8)


def _configure_numbering(doc: Document) -> None:
    try:
        root = doc.part.numbering_part.element
    except (AttributeError, NotImplementedError):
        return
    for level in root.xpath(".//w:lvl[@w:ilvl='0']"):
        ppr = level.find(qn("w:pPr"))
        if ppr is None:
            ppr = OxmlElement("w:pPr")
            level.append(ppr)
        ind = _ensure(ppr, "w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "279")
        tabs = _ensure(ppr, "w:tabs")
        tab = tabs.find(qn("w:tab"))
        if tab is None:
            tab = OxmlElement("w:tab")
            tabs.append(tab)
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")


def _format_figures_and_alt_text(doc: Document, *, language: str) -> None:
    descriptions = (
        [
            "Version-bound protocol flow: Client A owns matrix publication, "
            "metadata, the complete plan, and encrypted masks; Client B owns the "
            "query and key set, encrypts the gathered query, and reconstructs; "
            "the Cloud executes only the public program and encrypted operands.",
            "Evidence lineage from S1 and S2 through the stopped qualification "
            "to the Route C boundary; the formal campaign was not dispatched.",
        ]
        if language == "en"
        else [
            "版本绑定协议流程：Client A 拥有矩阵发布、元数据、完整计划和加密"
            "掩码；Client B 拥有查询与密钥、加密重排后的查询并完成重构；"
            "Cloud 只执行公开程序和密文操作数。",
            "从 S1、S2 到一次性资格停止点的证据链；正式实验没有启动，论文转入 Route C。",
        ]
    )
    image_index = 0
    for paragraph in doc.paragraphs:
        drawings = paragraph._p.xpath(".//w:drawing")
        if not drawings:
            continue
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.keep_with_next = True
        for drawing in drawings:
            for docpr in drawing.xpath(".//wp:docPr"):
                description = descriptions[min(image_index, len(descriptions) - 1)]
                docpr.set("descr", description)
                docpr.set("title", f"Figure {image_index + 1}")
            image_index += 1

    if image_index != 2:
        raise RuntimeError(f"expected exactly two manuscript figures, found {image_index}")


def _style_block_quotes(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        if paragraph.style.name in {"Block Text", "Quote"}:
            _paragraph_border_and_fill(paragraph, fill=LIGHT, left=BLUE)


def _apply_font_everywhere(doc: Document, font: str) -> None:
    for style in doc.styles:
        if (
            style.type in {WD_STYLE_TYPE.PARAGRAPH, WD_STYLE_TYPE.CHARACTER}
            and style.element.rPr is not None
        ):
            _set_font_element(style.element.get_or_add_rPr(), font)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            _set_run_font(run, font)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        _set_run_font(run, font)
    for section in doc.sections:
        for part in (
            section.header,
            section.first_page_header,
            section.footer,
            section.first_page_footer,
        ):
            for paragraph in part.paragraphs:
                for run in paragraph.runs:
                    _set_run_font(run, font)


def _left_align_long_identifier_paragraphs(doc: Document) -> None:
    """Avoid stretched word spacing around immutable IDs and provider digests."""
    body_styles = {"Normal", "Body Text", "First Paragraph"}
    for paragraph in doc.paragraphs:
        if paragraph.style.name not in body_styles:
            continue
        if any(len(token) >= 32 for token in paragraph.text.split()):
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _remove_empty_trailing_paragraphs(doc: Document) -> None:
    while (
        doc.paragraphs
        and not doc.paragraphs[-1].text
        and not doc.paragraphs[-1]._p.xpath(".//w:drawing")
    ):
        paragraph = doc.paragraphs[-1]
        paragraph._element.getparent().remove(paragraph._element)


def _drop_unreferenced_document_images(doc: Document) -> None:
    """Remove media inherited from a reference DOCX but unused by the body."""
    referenced = {
        relationship_id
        for element in doc.element.iter()
        for attribute in (qn("r:embed"), qn("r:link"))
        if (relationship_id := element.get(attribute)) is not None
    }
    for relationship_id, relationship in list(doc.part.rels.items()):
        if relationship.reltype == RT.IMAGE and relationship_id not in referenced:
            doc.part.drop_rel(relationship_id)


def _build_english(raw: Path, target: Path) -> None:
    doc = Document(raw)
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    doc.core_properties.title = EN_CORE_TITLE
    doc.core_properties.subject = "Route C methods and evidence-boundary working manuscript"
    doc.core_properties.keywords = (
        "homomorphic encryption; sparse matrix-vector multiplication; "
        "mutable sparse matrices; reproducible evaluation"
    )
    _configure_english_styles(doc)
    _promote_english_headings_and_add_kicker(doc)
    for section in doc.sections:
        _configure_page(section, letter=True)
        _set_header_footer(section, EN_SHORT, font="Calibri", first_label="")
    _configure_numbering(doc)
    _format_tables(doc, font="Calibri")
    _space_paragraphs_after_tables(doc)
    _format_figures_and_alt_text(doc, language="en")
    _style_block_quotes(doc)
    _apply_font_everywhere(doc, "Calibri")
    _left_align_long_identifier_paragraphs(doc)
    _remove_empty_trailing_paragraphs(doc)
    _drop_unreferenced_document_images(doc)
    doc.save(target)


def _build_chinese(raw: Path, target: Path) -> None:
    doc = Document(raw)
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    doc.core_properties.subject = "Route C 方法、证据边界与完整技术路线"
    doc.core_properties.keywords = "同态加密；稀疏矩阵向量乘法；可变稀疏矩阵；可复现评估"
    for section in doc.sections:
        _configure_page(section, letter=False)
        _set_header_footer(section, ZH_SHORT, font=ZH_FONT)
    _format_tables(doc, font=ZH_FONT)
    _space_paragraphs_after_tables(doc)
    _format_figures_and_alt_text(doc, language="zh")
    _apply_font_everywhere(doc, ZH_FONT)
    _remove_empty_trailing_paragraphs(doc)
    _drop_unreferenced_document_images(doc)
    doc.save(target)


def _validate_package(path: Path) -> None:
    with zipfile.ZipFile(path) as archive:
        bad = archive.testzip()
        if bad is not None:
            raise RuntimeError(f"corrupt DOCX member in {path}: {bad}")
        document_xml = archive.read("word/document.xml")
    if b"<m:oMath" not in document_xml:
        raise RuntimeError(f"editable OMML equations are absent from {path}")


def build(which: str) -> None:
    with tempfile.TemporaryDirectory(prefix="route-c-docx-") as temporary:
        temp = Path(temporary)
        if which in {"all", "english"}:
            raw = temp / "english.raw.docx"
            _pandoc(EN_MD, raw, reference=None, cite=True)
            staged = temp / "english.final.docx"
            _build_english(raw, staged)
            _validate_package(staged)
            shutil.copy2(staged, EN_DOCX)
        if which in {"all", "chinese"}:
            reference = temp / "chinese.reference.docx"
            shutil.copy2(ZH_DOCX, reference)
            raw = temp / "chinese.raw.docx"
            _pandoc(ZH_MD, raw, reference=reference, cite=False)
            staged = temp / "chinese.final.docx"
            _build_chinese(raw, staged)
            _validate_package(staged)
            shutil.copy2(staged, ZH_DOCX)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--which", choices=("all", "english", "chinese"), default="all")
    args = parser.parse_args()
    build(args.which)


if __name__ == "__main__":
    main()
